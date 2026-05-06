import io
from docx import Document
from fpdf import FPDF

def generate_docx(report_text, figure_name):
    doc = Document()
    doc.add_heading(f"Briefing d'Intelligence : {figure_name}", 0)
    
    blocks = report_text.split("---")
    for block in blocks:
        clean_block = block.strip()
        if clean_block:
            p = doc.add_paragraph()
            # Mini-parser to convert **text** into actual Bold text in Word
            parts = clean_block.split("**")
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0:  # Everything inside ** is an odd index
                    run.bold = True
            
            doc.add_paragraph("_" * 50) # Visual separator
            
    # Save to a virtual file in RAM (so we don't clutter your hard drive)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def generate_pdf(report_text, figure_name):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=25)
    
    # Title
    pdf.set_font("Helvetica", "B", 14)
    # Convert name to safe characters (PDFs without custom fonts struggle with Arabic)
    safe_name = figure_name.encode('latin-1', 'replace').decode('latin-1')
    pdf.cell(0, 10, f"Briefing d'Intelligence : {safe_name}", ln=True, align="C")
    pdf.ln(5)
    
    pdf.set_font("Helvetica", size=11)
    blocks = report_text.split("---")
    for block in blocks:
        clean_block = block.strip()
        if clean_block:
            # Remove Markdown bold tags for plain text PDF
            safe_block = clean_block.replace("**", "").encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 6, txt=safe_block)
            pdf.ln(5)
            pdf.cell(0, 0, "-" * 50, ln=True)
            pdf.ln(5)
            
    # Save to virtual file and return bytes
    output = pdf.output(dest="S")
    if isinstance(output, str):
        return output.encode("latin-1")
    return bytes(output)
